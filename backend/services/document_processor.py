"""
Document Processor - PRODUCTION VERSION
---------------------------------------
Real document processing using pypdf, pytesseract, and python-docx.
"""
import logging
import io
from typing import Dict, Any
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Processes documents using real libraries (not mocks).
    """
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._process_pdf,
            'image/png': self._process_image,
            'image/jpeg': self._process_image,
            'image/jpg': self._process_image,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._process_docx,
            'text/plain': self._process_text
        }
        logger.info("DocumentProcessor initialized (PRODUCTION MODE)")
    
    def process_file(self, file_bytes: bytes, filename: str, mime_type: str) -> Dict[str, Any]:
        """
        Process a document file using real libraries.
        
        Args:
            file_bytes: File content as bytes
            filename: Original filename
            mime_type: MIME type of the file
            
        Returns:
            Dictionary with extracted text and metadata
            
        Raises:
            ValueError: If file type not supported or processing fails
        """
        if not file_bytes:
            raise ValueError("File content cannot be empty")
        
        if mime_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {mime_type}")
        
        try:
            processor = self.supported_types[mime_type]
            result = processor(file_bytes, filename)
            result['processed_at'] = datetime.now().isoformat()
            result['filename'] = filename
            result['mime_type'] = mime_type
            result['size_bytes'] = len(file_bytes)
            
            logger.info(f"Processed {filename} ({mime_type}): {len(result.get('text', ''))} chars extracted")
            return result
        except Exception as e:
            logger.error(f"Document processing failed for {filename}: {e}", exc_info=True)
            raise ValueError(f"Failed to process document: {e}")
    
    def _process_pdf(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF using pypdf (or PyPDF2 compatibility fallback)."""
        try:
            # Prefer pypdf (actively maintained), but keep PyPDF2 fallback
            # for compatibility with legacy runtimes.
            try:
                from pypdf import PdfReader
            except ImportError:
                import PyPDF2  # type: ignore
                PdfReader = PyPDF2.PdfReader

            pdf_file = io.BytesIO(file_bytes)
            pdf_reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = "\n\n".join(text_parts)
            
            return {
                "text": full_text,
                "metadata": {
                    "type": "pdf",
                    "pages": len(pdf_reader.pages),
                    "encryption": pdf_reader.is_encrypted,
                    "metadata": pdf_reader.metadata if hasattr(pdf_reader, 'metadata') else {}
                }
            }
        except ImportError:
            logger.warning("pypdf not installed, using fallback")
            return {
                "text": "[PDF processing requires pypdf library]",
                "metadata": {"type": "pdf", "error": "pypdf not installed"}
            }
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            raise ValueError(f"PDF processing failed: {e}")
    
    def _process_image(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from image using pytesseract OCR."""
        try:
            from PIL import Image
            import pytesseract
            
            image = Image.open(io.BytesIO(file_bytes))
            text = pytesseract.image_to_string(image)
            
            return {
                "text": text,
                "metadata": {
                    "type": "image",
                    "format": image.format,
                    "size": image.size,
                    "mode": image.mode
                }
            }
        except ImportError as e:
            logger.warning(f"OCR libraries not installed: {e}")
            return {
                "text": "[Image OCR requires PIL and pytesseract libraries]",
                "metadata": {"type": "image", "error": "OCR libraries not installed"}
            }
        except Exception as e:
            logger.error(f"Image processing error: {e}")
            raise ValueError(f"Image OCR failed: {e}")
    
    def _process_docx(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from DOCX using python-docx."""
        try:
            from docx import Document
            
            doc = Document(io.BytesIO(file_bytes))
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]
            full_text = "\n".join(text_parts)
            
            return {
                "text": full_text,
                "metadata": {
                    "type": "docx",
                    "paragraphs": len(doc.paragraphs),
                    "sections": len(doc.sections) if hasattr(doc, 'sections') else 0
                }
            }
        except ImportError:
            logger.warning("python-docx not installed, using fallback")
            return {
                "text": "[DOCX processing requires python-docx library]",
                "metadata": {"type": "docx", "error": "python-docx not installed"}
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            raise ValueError(f"DOCX processing failed: {e}")
    
    def _process_text(self, file_bytes: bytes, filename: str) -> Dict[str, Any]:
        """Process plain text file."""
        try:
            text = file_bytes.decode('utf-8')
            return {
                "text": text,
                "metadata": {
                    "type": "text",
                    "encoding": "utf-8",
                    "lines": len(text.split('\n'))
                }
            }
        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_bytes.decode(encoding)
                    return {
                        "text": text,
                        "metadata": {
                            "type": "text",
                            "encoding": encoding,
                            "lines": len(text.split('\n'))
                        }
                    }
                except UnicodeDecodeError:
                    continue
            raise ValueError("Unable to decode text file with any common encoding")

# Global instance
document_processor = DocumentProcessor()
